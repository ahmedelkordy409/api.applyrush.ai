"""
Cover Letter export service for multiple formats.
"""

import io
import os
import uuid
from datetime import datetime, timezone, timedelta
from typing import Dict, Any, Optional
import structlog

# For PDF generation
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# For DOCX generation
try:
    from docx import Document
    from docx.shared import Inches
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

logger = structlog.get_logger(__name__)


class CoverLetterExportService:
    """Service for exporting cover letters in various formats."""

    def __init__(self, export_directory: str = "/tmp/cover_letter_exports"):
        self.export_directory = export_directory
        self.base_url = "https://api.applyrush.ai/exports"

        # Create export directory if it doesn't exist
        os.makedirs(export_directory, exist_ok=True)

    async def export_cover_letter(
        self,
        cover_letter_content: str,
        personal_info: Dict[str, str],
        export_format: str,
        include_contact_info: bool = True,
        letterhead_style: Optional[str] = None
    ) -> Dict[str, Any]:
        """Export cover letter in specified format."""
        try:
            export_id = str(uuid.uuid4())

            # Generate filename
            safe_name = personal_info.get("full_name", "Cover_Letter").replace(" ", "_")
            company_name = personal_info.get("company_name", "").replace(" ", "_")
            if company_name:
                filename = f"{safe_name}_{company_name}.{export_format}"
            else:
                filename = f"{safe_name}_{export_id[:8]}.{export_format}"

            file_path = os.path.join(self.export_directory, filename)

            # Export based on format
            if export_format.lower() == "pdf":
                file_size = await self._export_pdf(
                    cover_letter_content, personal_info, file_path,
                    include_contact_info, letterhead_style
                )
            elif export_format.lower() == "docx":
                file_size = await self._export_docx(
                    cover_letter_content, personal_info, file_path,
                    include_contact_info, letterhead_style
                )
            elif export_format.lower() == "txt":
                file_size = await self._export_txt(
                    cover_letter_content, personal_info, file_path,
                    include_contact_info
                )
            elif export_format.lower() == "html":
                file_size = await self._export_html(
                    cover_letter_content, personal_info, file_path,
                    include_contact_info, letterhead_style
                )
            else:
                raise ValueError(f"Unsupported export format: {export_format}")

            # Generate download URL (expires in 24 hours)
            download_url = f"{self.base_url}/{filename}"
            expires_at = datetime.now(timezone.utc) + timedelta(hours=24)

            return {
                "export_id": export_id,
                "download_url": download_url,
                "file_name": filename,
                "file_size_bytes": file_size,
                "format": export_format,
                "expires_at": expires_at,
                "file_path": file_path
            }

        except Exception as e:
            logger.error("Error exporting cover letter", error=str(e), format=export_format)
            raise

    async def _export_pdf(
        self,
        content: str,
        personal_info: Dict[str, str],
        file_path: str,
        include_contact_info: bool,
        letterhead_style: Optional[str]
    ) -> int:
        """Export as PDF format."""
        if not REPORTLAB_AVAILABLE:
            # Fallback to HTML if reportlab not available
            return await self._export_html(content, personal_info, file_path.replace('.pdf', '.html'), include_contact_info, letterhead_style)

        try:
            doc = SimpleDocTemplate(file_path, pagesize=letter, topMargin=1*inch)
            styles = getSampleStyleSheet()
            story = []

            # Header with contact info
            if include_contact_info:
                header_style = ParagraphStyle(
                    'HeaderStyle',
                    parent=styles['Normal'],
                    fontSize=12,
                    spaceAfter=20,
                    alignment=1  # Center alignment
                )

                header_text = f"""
                <b>{personal_info.get('full_name', '')}</b><br/>
                {personal_info.get('email_address', '')}<br/>
                {personal_info.get('phone_number', '')}<br/>
                {personal_info.get('city', '')}
                """
                story.append(Paragraph(header_text, header_style))
                story.append(Spacer(1, 0.2*inch))

            # Date
            date_style = ParagraphStyle(
                'DateStyle',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=20
            )
            current_date = datetime.now().strftime("%B %d, %Y")
            story.append(Paragraph(current_date, date_style))

            # Cover letter content
            content_style = ParagraphStyle(
                'ContentStyle',
                parent=styles['Normal'],
                fontSize=11,
                leading=14,
                spaceAfter=6
            )

            # Split content into paragraphs
            paragraphs = content.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    story.append(Paragraph(paragraph.strip(), content_style))
                    story.append(Spacer(1, 0.1*inch))

            doc.build(story)

            # Get file size
            return os.path.getsize(file_path)

        except Exception as e:
            logger.error("Error creating PDF", error=str(e))
            raise

    async def _export_docx(
        self,
        content: str,
        personal_info: Dict[str, str],
        file_path: str,
        include_contact_info: bool,
        letterhead_style: Optional[str]
    ) -> int:
        """Export as DOCX format."""
        if not PYTHON_DOCX_AVAILABLE:
            # Fallback to TXT if python-docx not available
            return await self._export_txt(content, personal_info, file_path.replace('.docx', '.txt'), include_contact_info)

        try:
            doc = Document()

            # Header with contact info
            if include_contact_info:
                header = doc.sections[0].header
                header_para = header.paragraphs[0]
                header_para.text = f"{personal_info.get('full_name', '')}\n" + \
                                  f"{personal_info.get('email_address', '')}\n" + \
                                  f"{personal_info.get('phone_number', '')}\n" + \
                                  f"{personal_info.get('city', '')}"

            # Date
            date_para = doc.add_paragraph()
            date_para.text = datetime.now().strftime("%B %d, %Y")

            # Add space
            doc.add_paragraph()

            # Cover letter content
            paragraphs = content.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())

            doc.save(file_path)

            # Get file size
            return os.path.getsize(file_path)

        except Exception as e:
            logger.error("Error creating DOCX", error=str(e))
            raise

    async def _export_txt(
        self,
        content: str,
        personal_info: Dict[str, str],
        file_path: str,
        include_contact_info: bool
    ) -> int:
        """Export as plain text format."""
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                # Header with contact info
                if include_contact_info:
                    f.write(f"{personal_info.get('full_name', '')}\n")
                    f.write(f"{personal_info.get('email_address', '')}\n")
                    f.write(f"{personal_info.get('phone_number', '')}\n")
                    f.write(f"{personal_info.get('city', '')}\n\n")

                # Date
                f.write(f"{datetime.now().strftime('%B %d, %Y')}\n\n")

                # Cover letter content
                f.write(content)

            # Get file size
            return os.path.getsize(file_path)

        except Exception as e:
            logger.error("Error creating TXT", error=str(e))
            raise

    async def _export_html(
        self,
        content: str,
        personal_info: Dict[str, str],
        file_path: str,
        include_contact_info: bool,
        letterhead_style: Optional[str]
    ) -> int:
        """Export as HTML format."""
        try:
            # Choose CSS style
            css_style = self._get_html_style(letterhead_style or "modern")

            html_content = f"""
            <!DOCTYPE html>
            <html lang="en">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>Cover Letter - {personal_info.get('full_name', '')}</title>
                <style>
                    {css_style}
                </style>
            </head>
            <body>
                <div class="container">
            """

            # Header with contact info
            if include_contact_info:
                html_content += f"""
                    <div class="header">
                        <h1>{personal_info.get('full_name', '')}</h1>
                        <div class="contact-info">
                            <p>{personal_info.get('email_address', '')}</p>
                            <p>{personal_info.get('phone_number', '')}</p>
                            <p>{personal_info.get('city', '')}</p>
                        </div>
                    </div>
                """

            # Date
            html_content += f"""
                <div class="date">
                    <p>{datetime.now().strftime('%B %d, %Y')}</p>
                </div>
            """

            # Cover letter content
            html_content += '<div class="content">'
            paragraphs = content.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    html_content += f'<p>{paragraph.strip()}</p>'

            html_content += """
                </div>
                </div>
            </body>
            </html>
            """

            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(html_content)

            # Get file size
            return os.path.getsize(file_path)

        except Exception as e:
            logger.error("Error creating HTML", error=str(e))
            raise

    def _get_html_style(self, style_name: str) -> str:
        """Get CSS styles for HTML export."""
        styles = {
            "modern": """
                body { font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; margin: 0; padding: 40px; background-color: #f8f9fa; }
                .container { max-width: 800px; margin: 0 auto; background: white; padding: 60px; box-shadow: 0 0 20px rgba(0,0,0,0.1); }
                .header { text-align: center; margin-bottom: 40px; border-bottom: 2px solid #007bff; padding-bottom: 20px; }
                .header h1 { margin: 0 0 20px 0; color: #007bff; font-size: 28px; font-weight: 600; }
                .contact-info p { margin: 5px 0; color: #6c757d; font-size: 14px; }
                .date { margin: 30px 0; font-size: 14px; color: #6c757d; }
                .content p { line-height: 1.6; margin-bottom: 16px; color: #333; font-size: 16px; }
            """,
            "classic": """
                body { font-family: 'Times New Roman', serif; margin: 0; padding: 40px; background-color: white; }
                .container { max-width: 800px; margin: 0 auto; }
                .header { text-align: center; margin-bottom: 40px; }
                .header h1 { margin: 0 0 20px 0; font-size: 24px; font-weight: bold; }
                .contact-info p { margin: 5px 0; font-size: 14px; }
                .date { margin: 30px 0; font-size: 14px; }
                .content p { line-height: 1.8; margin-bottom: 18px; font-size: 16px; text-align: justify; }
            """,
            "creative": """
                body { font-family: 'Georgia', serif; margin: 0; padding: 40px; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); }
                .container { max-width: 800px; margin: 0 auto; background: white; padding: 60px; border-radius: 15px; box-shadow: 0 20px 40px rgba(0,0,0,0.1); }
                .header { text-align: center; margin-bottom: 40px; }
                .header h1 { margin: 0 0 20px 0; color: #667eea; font-size: 32px; font-weight: 300; letter-spacing: 2px; }
                .contact-info p { margin: 8px 0; color: #555; font-size: 14px; font-style: italic; }
                .date { margin: 30px 0; font-size: 14px; color: #777; text-align: right; }
                .content p { line-height: 1.7; margin-bottom: 20px; color: #444; font-size: 16px; }
            """
        }
        return styles.get(style_name, styles["modern"])

    def cleanup_expired_files(self):
        """Clean up expired export files (run periodically)."""
        try:
            current_time = datetime.now()
            for filename in os.listdir(self.export_directory):
                file_path = os.path.join(self.export_directory, filename)
                if os.path.isfile(file_path):
                    # Check if file is older than 24 hours
                    file_time = datetime.fromtimestamp(os.path.getctime(file_path))
                    if (current_time - file_time).total_seconds() > 86400:  # 24 hours
                        os.remove(file_path)
                        logger.info("Cleaned up expired export file", filename=filename)

        except Exception as e:
            logger.error("Error cleaning up export files", error=str(e))